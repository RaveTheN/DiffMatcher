#!/usr/bin/env python3
"""
Test script for DiffMatcher
Creates test files and runs basic functionality tests
Includes Word document testing if python-docx is available
"""

import os
import tempfile
from pathlib import Path
import sys

# Add current directory to path for imports
sys.path.append(str(Path(__file__).parent))

from cli_diff_matcher import compare_files_line_by_line, extract_text_from_file

# Check for Word document support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def create_test_files():
    """Create temporary test files"""
    # Create temporary directory
    temp_dir = Path(tempfile.mkdtemp())
    
    # Test file 1
    file1_content = """Line 1: Identical
Line 2: Identical  
Line 3: Different in file 1
Line 4: Identical
Line 5: Only in file 1"""
    
    # Test file 2
    file2_content = """Line 1: Identical
Line 2: Identical
Line 3: Different in file 2
Line 4: Identical
Line 6: Only in file 2
Line 7: Additional line"""
    
    file1_path = temp_dir / "test_file_1.txt"
    file2_path = temp_dir / "test_file_2.txt"
    
    with open(file1_path, 'w', encoding='utf-8') as f:
        f.write(file1_content)
    
    with open(file2_path, 'w', encoding='utf-8') as f:
        f.write(file2_content)
    
    return str(file1_path), str(file2_path), temp_dir


def test_identical_files():
    """Test with identical files"""
    print("🧪 Testing identical files...")
    
    temp_dir = Path(tempfile.mkdtemp())
    content = "Line 1\nLine 2\nLine 3\n"
    
    file1 = temp_dir / "identical1.txt"
    file2 = temp_dir / "identical2.txt"
    
    with open(file1, 'w') as f:
        f.write(content)
    with open(file2, 'w') as f:
        f.write(content)
    
    similarity = compare_files_line_by_line(str(file1), str(file2), verbose=False)
    
    assert similarity == 100.0, f"Expected 100% similarity, got {similarity}%"
    print("✅ Identical files test passed")
    
    # Cleanup
    os.unlink(file1)
    os.unlink(file2)
    os.rmdir(temp_dir)


def test_different_files():
    """Test with different files"""
    print("🧪 Testing different files...")
    
    file1, file2, temp_dir = create_test_files()
    
    similarity = compare_files_line_by_line(file1, file2, verbose=False)
    
    # Should have some similarity but not 100%
    assert 0 < similarity < 100, f"Expected partial similarity, got {similarity}%"
    print(f"✅ Different files test passed (similarity: {similarity}%)")
    
    # Cleanup
    os.unlink(file1)
    os.unlink(file2)
    os.rmdir(temp_dir)


def test_empty_files():
    """Test with empty files"""
    print("🧪 Testing empty files...")
    
    temp_dir = Path(tempfile.mkdtemp())
    
    file1 = temp_dir / "empty1.txt"
    file2 = temp_dir / "empty2.txt"
    
    with open(file1, 'w') as f:
        pass  # Create empty file
    with open(file2, 'w') as f:
        pass  # Create empty file
    
    similarity = compare_files_line_by_line(str(file1), str(file2), verbose=False)
    
    assert similarity == 0.0, f"Expected 0% similarity for empty files, got {similarity}%"
    print("✅ Empty files test passed")
    
    # Cleanup
    os.unlink(file1)
    os.unlink(file2)
    os.rmdir(temp_dir)


def test_gui_import():
    """Test that GUI module can be imported"""
    print("🧪 Testing GUI module import...")
    
    try:
        import diff_matcher
        print("✅ GUI module import test passed")
        return True
    except ImportError as e:
        print(f"❌ GUI module import test failed: {e}")
        return False


def test_word_documents():
    """Test with Word documents if python-docx is available"""
    if not DOCX_AVAILABLE:
        print("🧪 Skipping Word document test - python-docx not available")
        return True
    
    print("🧪 Testing Word documents...")
    
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        # Create Word document 1
        doc1_path = temp_dir / "test_doc1.docx"
        doc1 = Document()
        doc1.add_paragraph("Line 1: Identical")
        doc1.add_paragraph("Line 2: Different in doc 1")
        doc1.add_paragraph("Line 3: Identical")
        doc1.save(doc1_path)
        
        # Create Word document 2
        doc2_path = temp_dir / "test_doc2.docx"
        doc2 = Document()
        doc2.add_paragraph("Line 1: Identical")
        doc2.add_paragraph("Line 2: Different in doc 2")
        doc2.add_paragraph("Line 3: Identical")
        doc2.save(doc2_path)
        
        similarity = compare_files_line_by_line(str(doc1_path), str(doc2_path), verbose=False)
        
        # Should have some similarity but not 100%
        assert 0 < similarity < 100, f"Expected partial similarity, got {similarity}%"
        print(f"✅ Word documents test passed (similarity: {similarity}%)")
        
        return True
        
    except Exception as e:
        print(f"❌ Word documents test failed: {e}")
        return False
    
    finally:
        # Cleanup
        for file in temp_dir.glob("*"):
            try:
                os.unlink(file)
            except:
                pass
        try:
            os.rmdir(temp_dir)
        except:
            pass


def test_mixed_file_types():
    """Test comparing text file with Word document"""
    if not DOCX_AVAILABLE:
        print("🧪 Skipping mixed file types test - python-docx not available")
        return True
    
    print("🧪 Testing mixed file types (text vs Word)...")
    
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        # Create text file
        txt_path = temp_dir / "test.txt"
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write("Line 1\nLine 2\nLine 3\n")
        
        # Create Word document with same content
        docx_path = temp_dir / "test.docx"
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3")
        doc.save(docx_path)
        
        similarity = compare_files_line_by_line(str(txt_path), str(docx_path), verbose=False)
        
        # Should be highly similar (near 100%)
        assert similarity > 90, f"Expected high similarity, got {similarity}%"
        print(f"✅ Mixed file types test passed (similarity: {similarity}%)")
        
        return True
        
    except Exception as e:
        print(f"❌ Mixed file types test failed: {e}")
        return False
    
    finally:
        # Cleanup
        for file in temp_dir.glob("*"):
            try:
                os.unlink(file)
            except:
                pass
        try:
            os.rmdir(temp_dir)
        except:
            pass


def run_all_tests():
    """Run all tests"""
    print("🚀 Running DiffMatcher Tests (including Word document support)")
    print("=" * 60)
    
    tests_passed = 0
    total_tests = 6  # Updated total
    
    try:
        test_identical_files()
        tests_passed += 1
    except Exception as e:
        print(f"❌ Identical files test failed: {e}")
    
    try:
        test_different_files()
        tests_passed += 1
    except Exception as e:
        print(f"❌ Different files test failed: {e}")
    
    try:
        test_empty_files()
        tests_passed += 1
    except Exception as e:
        print(f"❌ Empty files test failed: {e}")
    
    try:
        if test_gui_import():
            tests_passed += 1
    except Exception as e:
        print(f"❌ GUI import test failed: {e}")
    
    try:
        if test_word_documents():
            tests_passed += 1
    except Exception as e:
        print(f"❌ Word documents test failed: {e}")
    
    try:
        if test_mixed_file_types():
            tests_passed += 1
    except Exception as e:
        print(f"❌ Mixed file types test failed: {e}")
    
    print("\n" + "=" * 60)
    print(f"📊 Test Results: {tests_passed}/{total_tests} tests passed")
    
    if DOCX_AVAILABLE:
        print("📄 Word document support: ✅ Available and tested")
    else:
        print("📄 Word document support: ❌ Not available (python-docx not installed)")
    
    if tests_passed == total_tests:
        print("🎉 All tests passed!")
        return True
    else:
        print("⚠️ Some tests failed")
        return False


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
