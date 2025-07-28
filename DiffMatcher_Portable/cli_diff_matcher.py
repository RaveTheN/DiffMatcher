#!/usr/bin/env python3
"""
DiffMatcher CLI - Command Line File Comparison Tool
Based on the original line-by-line comparison function
Supports text files and Microsoft Word documents (.docx)
"""

import argparse
import sys
from pathlib import Path
from difflib import SequenceMatcher

# Import for Word document support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def extract_text_from_file(file_path):
    """
    Extract text content from different file types
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        list: List of lines from the file
    """
    file_path = Path(file_path)
    
    try:
        if file_path.suffix.lower() == '.docx':
            if not DOCX_AVAILABLE:
                raise Exception("python-docx library is not installed. Cannot read .docx files.\nInstall with: pip install python-docx")
            
            # Extract text from Word document
            doc = Document(file_path)
            lines = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                lines.append(paragraph.text + '\n')
            
            # If no paragraphs found, try tables
            if not lines:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                lines.append(cell.text + '\n')
            
            return lines
        
        else:
            # Handle text files (including .txt, .py, etc.)
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.readlines()
                
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.readlines()
        except:
            raise Exception(f"Could not decode file: {file_path}")
    except Exception as e:
        raise Exception(f"Error reading file {file_path}: {str(e)}")


def compare_files_line_by_line(file1, file2, verbose=True):
    """
    Compare two files line by line and return similarity percentage
    Supports text files and Word documents (.docx)
    
    Args:
        file1 (str): Path to first file
        file2 (str): Path to second file
        verbose (bool): Whether to print detailed differences
    
    Returns:
        float: Average similarity percentage
    """
    try:
        lines1 = extract_text_from_file(file1)
        lines2 = extract_text_from_file(file2)
    except FileNotFoundError as e:
        print(f"❌ Error: File not found - {e}")
        return None
    except Exception as e:
        print(f"❌ Error reading files: {e}")
        return None

    total_similarity = 0
    lines_compared = 0
    differences_count = 0

    max_lines = max(len(lines1), len(lines2))

    # Determine file types for display
    file1_type = "Word document" if Path(file1).suffix.lower() == '.docx' else "Text file"
    file2_type = "Word document" if Path(file2).suffix.lower() == '.docx' else "Text file"

    print(f"\n📊 COMPARISON SUMMARY:")
    print(f"   File 1: {Path(file1).name} ({file1_type}, {len(lines1)} lines)")
    print(f"   File 2: {Path(file2).name} ({file2_type}, {len(lines2)} lines)")
    print(f"   Total lines to compare: {max_lines}")

    for i in range(max_lines):
        line1 = lines1[i].strip() if i < len(lines1) else ''
        line2 = lines2[i].strip() if i < len(lines2) else ''
        similarity = SequenceMatcher(None, line1, line2).ratio()
        total_similarity += similarity
        lines_compared += 1

        if similarity < 1.0:
            differences_count += 1
            if verbose:
                print(f"\n🛑 Line {i + 1} differs:")
                print(f"   File 1: {line1 if line1 else '(empty line)'}")
                print(f"   File 2: {line2 if line2 else '(empty line)'}")
                print(f"   Similarity: {round(similarity * 100, 2)}%")

    if lines_compared == 0:
        print("⚠️ No lines to compare")
        return 0.0

    avg_similarity = (total_similarity / lines_compared) * 100
    
    print(f"\n📈 RESULTS:")
    print(f"   Differences found: {differences_count}")
    print(f"   Average similarity: {round(avg_similarity, 2)}%")
    
    # Provide interpretation
    if avg_similarity >= 95:
        print("   ✅ Files are nearly identical!")
    elif avg_similarity >= 80:
        print("   ⚠️ Files are quite similar with some differences")
    elif avg_similarity >= 50:
        print("   🔶 Files have moderate similarity")
    else:
        print("   ❌ Files are significantly different")
    
    return round(avg_similarity, 2)


def create_sample_files():
    """Create sample files for testing"""
    sample_dir = Path("sample_files")
    sample_dir.mkdir(exist_ok=True)
    
    # Sample file 1
    sample1_content = """Hello World!
This is a sample file for testing.
It contains multiple lines of text.
Some lines will be identical.
Others will be slightly different.
This line exists in both files.
End of file 1."""

    # Sample file 2
    sample2_content = """Hello World!
This is a sample file for testing purposes.
It contains multiple lines of text.
Some lines will be identical.
Others will be very different.
This line exists in both files.
Additional line in file 2.
End of file 2."""
    
    # Create text files
    file1_path = sample_dir / "sample_file_1.txt"
    file2_path = sample_dir / "sample_file_2.txt"
    
    with open(file1_path, 'w', encoding='utf-8') as f:
        f.write(sample1_content)
    
    with open(file2_path, 'w', encoding='utf-8') as f:
        f.write(sample2_content)
    
    # Create Word documents if python-docx is available
    docx1_path = None
    docx2_path = None
    
    if DOCX_AVAILABLE:
        try:
            # Create Word document 1
            docx1_path = sample_dir / "sample_document_1.docx"
            doc1 = Document()
            for line in sample1_content.split('\n'):
                if line.strip():
                    doc1.add_paragraph(line)
            doc1.save(docx1_path)
            
            # Create Word document 2
            docx2_path = sample_dir / "sample_document_2.docx"
            doc2 = Document()
            for line in sample2_content.split('\n'):
                if line.strip():
                    doc2.add_paragraph(line)
            doc2.save(docx2_path)
            
        except Exception as e:
            print(f"Warning: Could not create Word documents: {e}")
    
    print(f"✅ Sample files created:")
    print(f"   📁 {file1_path}")
    print(f"   📁 {file2_path}")
    
    if DOCX_AVAILABLE and docx1_path and docx2_path:
        print(f"   📄 {docx1_path}")
        print(f"   📄 {docx2_path}")
        return str(docx1_path), str(docx2_path)
    else:
        if not DOCX_AVAILABLE:
            print("   ℹ️ Install python-docx to create Word document samples")
        return str(file1_path), str(file2_path)


def main():
    """Main function for CLI interface"""
    parser = argparse.ArgumentParser(
        description="DiffMatcher CLI - Compare two files line by line (supports .txt and .docx)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python cli_diff_matcher.py file1.txt file2.txt
  python cli_diff_matcher.py document1.docx document2.docx
  python cli_diff_matcher.py file1.txt document2.docx --quiet
  python cli_diff_matcher.py --sample

Supported file types:
  • Text files (.txt, .py, etc.)
  • Microsoft Word documents (.docx) - requires python-docx
        """
    )
    
    parser.add_argument('file1', nargs='?', help='First file to compare')
    parser.add_argument('file2', nargs='?', help='Second file to compare')
    parser.add_argument('--quiet', '-q', action='store_true', 
                       help='Suppress detailed output, show only summary')
    parser.add_argument('--sample', '-s', action='store_true',
                       help='Create sample files and compare them')
    parser.add_argument('--version', '-v', action='version', version='DiffMatcher CLI 2.0 (with Word support)')
    
    args = parser.parse_args()
    
    print("🔍 DiffMatcher CLI - File Comparison Tool")
    if DOCX_AVAILABLE:
        print("📄 Word document support: ✅ Enabled")
    else:
        print("📄 Word document support: ❌ Disabled (install python-docx)")
    print("=" * 50)
    
    if args.sample:
        print("📝 Creating sample files...")
        file1, file2 = create_sample_files()
        print(f"\n🚀 Comparing sample files...")
        similarity = compare_files_line_by_line(file1, file2, verbose=not args.quiet)
        
    elif args.file1 and args.file2:
        # Validate files exist
        if not Path(args.file1).exists():
            print(f"❌ Error: File '{args.file1}' does not exist")
            sys.exit(1)
        
        if not Path(args.file2).exists():
            print(f"❌ Error: File '{args.file2}' does not exist")
            sys.exit(1)
        
        # Check if Word documents are being used without python-docx
        file1_is_docx = Path(args.file1).suffix.lower() == '.docx'
        file2_is_docx = Path(args.file2).suffix.lower() == '.docx'
        
        if (file1_is_docx or file2_is_docx) and not DOCX_AVAILABLE:
            print("❌ Error: Word document detected but python-docx is not installed")
            print("   Install with: pip install python-docx")
            sys.exit(1)
        
        print(f"🚀 Comparing files...")
        similarity = compare_files_line_by_line(args.file1, args.file2, verbose=not args.quiet)
        
    else:
        print("❌ Error: Please provide two files to compare or use --sample flag")
        parser.print_help()
        sys.exit(1)
    
    if similarity is not None:
        print(f"\n🎯 Final Result: {similarity}% similarity")
        
        # Exit code based on similarity
        if similarity >= 95:
            sys.exit(0)  # Files are nearly identical
        elif similarity >= 50:
            sys.exit(1)  # Files have differences but are similar
        else:
            sys.exit(2)  # Files are significantly different
    else:
        sys.exit(3)  # Error occurred


if __name__ == "__main__":
    main()
