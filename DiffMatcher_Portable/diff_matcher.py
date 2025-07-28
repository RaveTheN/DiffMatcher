import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
from difflib import SequenceMatcher
import os
from pathlib import Path

# Import for Word document support
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DiffMatcher:
    def __init__(self, root):
        self.root = root
        self.root.title("DiffMatcher - File Comparison Tool")
        self.root.geometry("900x700")
        self.root.configure(bg="#f0f0f0")
        
        # Variables to store file paths
        self.file1_path = tk.StringVar()
        self.file2_path = tk.StringVar()
        
        self.create_widgets()
    
    def create_widgets(self):
        """Create and arrange all GUI widgets"""
        # Main frame
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(4, weight=1)
        
        # Title
        title_label = tk.Label(main_frame, text="DiffMatcher", 
                              font=("Arial", 16, "bold"), 
                              bg="#f0f0f0", fg="#333")
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 10))
        
        # Word support status
        word_support_text = "📄 Word document support: ✅ Enabled" if DOCX_AVAILABLE else "📄 Word document support: ❌ Disabled (install python-docx)"
        word_status_label = tk.Label(main_frame, text=word_support_text,
                                    font=("Arial", 9),
                                    bg="#f0f0f0", fg="#666")
        word_status_label.grid(row=0, column=0, columnspan=3, pady=(25, 20), sticky=tk.S)
        
        # File 1 selection
        ttk.Label(main_frame, text="File 1:").grid(row=1, column=0, sticky=tk.W, pady=5)
        file1_entry = ttk.Entry(main_frame, textvariable=self.file1_path, width=60)
        file1_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=5, padx=(5, 5))
        ttk.Button(main_frame, text="Browse", 
                  command=lambda: self.browse_file(self.file1_path)).grid(row=1, column=2, pady=5)
        
        # File 2 selection
        ttk.Label(main_frame, text="File 2:").grid(row=2, column=0, sticky=tk.W, pady=5)
        file2_entry = ttk.Entry(main_frame, textvariable=self.file2_path, width=60)
        file2_entry.grid(row=2, column=1, sticky=(tk.W, tk.E), pady=5, padx=(5, 5))
        ttk.Button(main_frame, text="Browse", 
                  command=lambda: self.browse_file(self.file2_path)).grid(row=2, column=2, pady=5)
        
        # Buttons frame
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.grid(row=3, column=0, columnspan=3, pady=20)
        
        # Compare button
        compare_btn = ttk.Button(buttons_frame, text="Compare Files", 
                               command=self.compare_files, style="Accent.TButton")
        compare_btn.pack(side=tk.LEFT, padx=5)
        
        # Clear button
        clear_btn = ttk.Button(buttons_frame, text="Clear Results", 
                             command=self.clear_results)
        clear_btn.pack(side=tk.LEFT, padx=5)
        
        # Create sample files button
        sample_btn = ttk.Button(buttons_frame, text="Create Sample Files", 
                              command=self.create_sample_files)
        sample_btn.pack(side=tk.LEFT, padx=5)
        
        # Results area
        results_frame = ttk.LabelFrame(main_frame, text="Comparison Results", padding="10")
        results_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(10, 0))
        results_frame.columnconfigure(0, weight=1)
        results_frame.rowconfigure(0, weight=1)
        
        # Text widget for results
        self.results_text = scrolledtext.ScrolledText(results_frame, 
                                                     wrap=tk.WORD, 
                                                     height=20,
                                                     font=("Consolas", 10))
        self.results_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(10, 0))
        
        # Status bar
        self.status_var = tk.StringVar(value="Ready to compare files")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, 
                              relief=tk.SUNKEN, anchor=tk.W)
        status_bar.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(5, 0))
    
    def browse_file(self, file_var):
        """Open file dialog to select a file"""
        filetypes = [
            ("All supported files", "*.txt;*.py;*.docx"),
            ("Text files", "*.txt"),
            ("Python files", "*.py"),
        ]
        
        # Add Word documents if python-docx is available
        if DOCX_AVAILABLE:
            filetypes.insert(1, ("Word documents", "*.docx"))
        else:
            # Show warning about Word document support
            filetypes.insert(1, ("Word documents (⚠️ requires python-docx)", "*.docx"))
        
        # Always add "All files" at the end
        filetypes.append(("All files", "*.*"))
        
        filename = filedialog.askopenfilename(
            title="Select a file to compare",
            filetypes=filetypes
        )
        if filename:
            # Check if Word document is selected without support
            if filename.lower().endswith('.docx') and not DOCX_AVAILABLE:
                response = messagebox.askyesno(
                    "Word Document Support Required",
                    "You selected a Word document (.docx) but python-docx is not installed.\n\n"
                    "Word documents cannot be compared without this library.\n\n"
                    "Do you want to continue anyway? (The comparison will fail)",
                    icon="warning"
                )
                if not response:
                    return  # Don't set the file path
            
            file_var.set(filename)
    
    def extract_text_from_file(self, file_path):
        """
        Extract text content from different file types
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            list: List of lines from the file
        """
        file_path = Path(file_path)
        
        try:
            if file_path.suffix.lower() == '.docx':
                if not DOCX_AVAILABLE:
                    # Show user-friendly error message
                    error_msg = (
                        "Word document (.docx) detected but python-docx library is not installed.\n\n"
                        "To enable Word document support:\n"
                        "1. Open a terminal/command prompt\n"
                        "2. Run: pip install python-docx\n"
                        "3. Restart the application\n\n"
                        "Alternative: Use the virtual environment:\n"
                        "Run the GUI with: C:/Progetti/DiffMatcher/.venv/Scripts/python.exe diff_matcher.py"
                    )
                    raise Exception(error_msg)
                
                # Extract text from Word document
                doc = Document(file_path)
                lines = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    lines.append(paragraph.text + '\n')
                
                # If no paragraphs found, try tables
                if not lines:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    lines.append(cell.text + '\n')
                
                return lines
            
            else:
                # Handle text files (including .txt, .py, etc.)
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.readlines()
                    
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.readlines()
            except:
                raise Exception(f"Could not decode file: {file_path}")
        except Exception as e:
            raise Exception(f"Error reading file {file_path}: {str(e)}")
    
    def compare_files_line_by_line(self, file1, file2):
        """
        Compare two files line by line and return detailed results
        Based on the original function with GUI integration and multi-format support
        """
        try:
            lines1 = self.extract_text_from_file(file1)
            lines2 = self.extract_text_from_file(file2)
        except Exception as e:
            raise Exception(f"Error reading files: {str(e)}")

        total_similarity = 0
        lines_compared = 0
        differences = []

        max_lines = max(len(lines1), len(lines2))

        for i in range(max_lines):
            line1 = lines1[i].strip() if i < len(lines1) else ''
            line2 = lines2[i].strip() if i < len(lines2) else ''
            similarity = SequenceMatcher(None, line1, line2).ratio()
            total_similarity += similarity
            lines_compared += 1

            if similarity < 1.0:
                differences.append({
                    'line_num': i + 1,
                    'line1': line1,
                    'line2': line2,
                    'similarity': round(similarity * 100, 2)
                })

        if lines_compared == 0:
            return 0.0, differences, 0, 0

        avg_similarity = (total_similarity / lines_compared) * 100
        return round(avg_similarity, 2), differences, len(lines1), len(lines2)
    
    def compare_files(self):
        """Compare the selected files and display results"""
        file1 = self.file1_path.get()
        file2 = self.file2_path.get()
        
        # Validation
        if not file1 or not file2:
            messagebox.showerror("Error", "Please select both files to compare")
            return
        
        if not os.path.exists(file1):
            messagebox.showerror("Error", f"File 1 does not exist: {file1}")
            return
        
        if not os.path.exists(file2):
            messagebox.showerror("Error", f"File 2 does not exist: {file2}")
            return
        
        # Start progress bar
        self.progress.start()
        self.status_var.set("Comparing files...")
        self.root.update()
        
        try:
            # Perform comparison
            similarity, differences, lines1_count, lines2_count = self.compare_files_line_by_line(file1, file2)
            
            # Display results
            self.display_results(file1, file2, similarity, differences, lines1_count, lines2_count)
            
            self.status_var.set(f"Comparison complete - {similarity}% similarity")
            
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred during comparison:\n{str(e)}")
            self.status_var.set("Error occurred during comparison")
        
        finally:
            # Stop progress bar
            self.progress.stop()
    
    def display_results(self, file1, file2, similarity, differences, lines1_count, lines2_count):
        """Display comparison results in the text widget"""
        self.results_text.delete(1.0, tk.END)
        
        # Header
        result_text = f"🔍 FILE COMPARISON RESULTS\n"
        result_text += f"=" * 50 + "\n\n"
        result_text += f"📁 File 1: {Path(file1).name}\n"
        result_text += f"📁 File 2: {Path(file2).name}\n\n"
        result_text += f"📊 SUMMARY:\n"
        result_text += f"   • File 1 lines: {lines1_count}\n"
        result_text += f"   • File 2 lines: {lines2_count}\n"
        result_text += f"   • Average similarity: {similarity}%\n"
        result_text += f"   • Differences found: {len(differences)}\n\n"
        
        # Overall assessment
        if similarity >= 95:
            result_text += f"✅ Files are nearly identical!\n\n"
        elif similarity >= 80:
            result_text += f"⚠️ Files are quite similar with some differences.\n\n"
        elif similarity >= 50:
            result_text += f"🔶 Files have moderate similarity.\n\n"
        else:
            result_text += f"❌ Files are significantly different.\n\n"
        
        # Detailed differences
        if differences:
            result_text += f"🛑 DETAILED DIFFERENCES:\n"
            result_text += f"-" * 30 + "\n\n"
            
            for diff in differences:  # Show all differences
                result_text += f"Line {diff['line_num']} (Similarity: {diff['similarity']}%):\n"
                result_text += f"  File 1: {diff['line1'] or '(empty line)'}\n"
                result_text += f"  File 2: {diff['line2'] or '(empty line)'}\n\n"
        else:
            result_text += f"✨ No differences found - files are identical!\n"
        
        self.results_text.insert(1.0, result_text)
    
    def clear_results(self):
        """Clear the results area"""
        self.results_text.delete(1.0, tk.END)
        self.file1_path.set("")
        self.file2_path.set("")
        self.status_var.set("Ready to compare files")
    
    def create_sample_files(self):
        """Create sample files for testing"""
        try:
            # Sample file 1
            sample1_content = """Hello World!
This is a sample file for testing.
It contains multiple lines of text.
Some lines will be identical.
Others will be slightly different.
This line exists in both files.
End of file 1."""

            # Sample file 2
            sample2_content = """Hello World!
This is a sample file for testing purposes.
It contains multiple lines of text.
Some lines will be identical.
Others will be very different.
This line exists in both files.
Additional line in file 2.
End of file 2."""
            
            # Create sample directory if it doesn't exist
            sample_dir = Path(os.getcwd()) / "sample_files"
            sample_dir.mkdir(exist_ok=True)
            
            # Create text files
            file1_path = sample_dir / "sample_file_1.txt"
            file2_path = sample_dir / "sample_file_2.txt"
            
            # Write sample text files
            with open(file1_path, 'w', encoding='utf-8') as f:
                f.write(sample1_content)
            
            with open(file2_path, 'w', encoding='utf-8') as f:
                f.write(sample2_content)
            
            # Create Word documents if python-docx is available
            docx1_path = None
            docx2_path = None
            
            if DOCX_AVAILABLE:
                try:
                    # Create Word document 1
                    docx1_path = sample_dir / "sample_document_1.docx"
                    doc1 = Document()
                    for line in sample1_content.split('\n'):
                        if line.strip():
                            doc1.add_paragraph(line)
                    doc1.save(docx1_path)
                    
                    # Create Word document 2
                    docx2_path = sample_dir / "sample_document_2.docx"
                    doc2 = Document()
                    for line in sample2_content.split('\n'):
                        if line.strip():
                            doc2.add_paragraph(line)
                    doc2.save(docx2_path)
                    
                except Exception as e:
                    print(f"Warning: Could not create Word documents: {e}")
            
            # Set file paths in the GUI (prefer Word documents if available)
            if DOCX_AVAILABLE and docx1_path and docx2_path:
                self.file1_path.set(str(docx1_path))
                self.file2_path.set(str(docx2_path))
                file_type = "Word documents and text files"
            else:
                self.file1_path.set(str(file1_path))
                self.file2_path.set(str(file2_path))
                file_type = "text files"
            
            success_msg = f"Sample {file_type} created successfully!\n\nFiles created in: {sample_dir}"
            if DOCX_AVAILABLE:
                success_msg += "\n\nBoth .txt and .docx versions created."
            else:
                success_msg += "\n\nNote: Install python-docx to create Word document samples."
                
            messagebox.showinfo("Success", success_msg)
            self.status_var.set("Sample files created and loaded")
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create sample files:\n{str(e)}")


def main():
    """Main function to run the application"""
    root = tk.Tk()
    app = DiffMatcher(root)
    root.mainloop()


if __name__ == "__main__":
    main()
