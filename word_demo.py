#!/usr/bin/env python3
"""
Word Document Demo for DiffMatcher
Creates sample Word documents with different content to demonstrate comparison
"""

import sys
from pathlib import Path

# Add current directory to path for imports
sys.path.append(str(Path(__file__).parent))

try:
    from docx import Document
    from cli_diff_matcher import compare_files_line_by_line
    DOCX_AVAILABLE = True
except ImportError:
    print("❌ python-docx not available. Install with: pip install python-docx")
    sys.exit(1)


def create_word_demo():
    """Create demonstration Word documents with complex content"""
    print("📄 Creating Word Document Comparison Demo")
    print("=" * 50)
    
    sample_dir = Path("sample_files")
    sample_dir.mkdir(exist_ok=True)
    
    # Document 1: Original report
    doc1_path = sample_dir / "report_v1.docx"
    doc1 = Document()
    
    # Add title
    title1 = doc1.add_heading('Project Status Report - Version 1.0', 0)
    
    # Add paragraphs
    doc1.add_paragraph('Executive Summary')
    doc1.add_paragraph('This report covers the current status of our software development project.')
    doc1.add_paragraph('The project is currently on track and within budget.')
    doc1.add_paragraph('')
    doc1.add_paragraph('Key Achievements:')
    doc1.add_paragraph('• Completed user interface design')
    doc1.add_paragraph('• Implemented core functionality')
    doc1.add_paragraph('• Conducted initial testing')
    doc1.add_paragraph('')
    doc1.add_paragraph('Next Steps:')
    doc1.add_paragraph('• Finalize testing procedures')
    doc1.add_paragraph('• Prepare for deployment')
    doc1.add_paragraph('• Document user manual')
    
    doc1.save(doc1_path)
    
    # Document 2: Updated report with changes
    doc2_path = sample_dir / "report_v2.docx"
    doc2 = Document()
    
    # Add title (slightly different)
    title2 = doc2.add_heading('Project Status Report - Version 2.0', 0)
    
    # Add paragraphs with modifications
    doc2.add_paragraph('Executive Summary')
    doc2.add_paragraph('This report covers the current status of our software development project.')
    doc2.add_paragraph('The project is slightly behind schedule but remains within budget.')  # Changed
    doc2.add_paragraph('')
    doc2.add_paragraph('Key Achievements:')
    doc2.add_paragraph('• Completed user interface design')
    doc2.add_paragraph('• Implemented core functionality')
    doc2.add_paragraph('• Conducted comprehensive testing')  # Changed
    doc2.add_paragraph('• Fixed critical bugs')  # New line
    doc2.add_paragraph('')
    doc2.add_paragraph('Next Steps:')
    doc2.add_paragraph('• Complete final testing procedures')  # Changed
    doc2.add_paragraph('• Prepare for deployment')
    doc2.add_paragraph('• Document user manual')
    doc2.add_paragraph('• Train support team')  # New line
    
    doc2.save(doc2_path)
    
    print(f"✅ Created Word documents:")
    print(f"   📄 {doc1_path}")
    print(f"   📄 {doc2_path}")
    
    return str(doc1_path), str(doc2_path)


def main():
    """Main demo function"""
    print("🎯 DiffMatcher Word Document Demo")
    print("=" * 50)
    
    # Create demo documents
    doc1, doc2 = create_word_demo()
    
    print(f"\n🚀 Comparing Word documents...")
    print(f"   File 1: {Path(doc1).name}")
    print(f"   File 2: {Path(doc2).name}")
    
    # Compare the documents
    similarity = compare_files_line_by_line(doc1, doc2, verbose=True)
    
    print(f"\n" + "=" * 50)
    print(f"🎯 Final Comparison Result: {similarity}% similarity")
    print(f"📊 This demonstrates DiffMatcher's ability to:")
    print(f"   • Extract text from Word documents")
    print(f"   • Compare document structure and content")
    print(f"   • Identify specific differences between versions")
    print(f"   • Provide detailed similarity analysis")
    
    print(f"\n💡 You can now use these files to test the GUI:")
    print(f"   python diff_matcher.py")
    print(f"   Then browse and select: {doc1} and {doc2}")


if __name__ == "__main__":
    main()
